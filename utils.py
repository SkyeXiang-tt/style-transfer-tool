"""
utils.py — 辅助函数模块
包含：AI 调用、文件读写、风格树管理、Markdown→Word 转换、去 AI 化
"""

from pathlib import Path
from datetime import datetime
import re
import io
import shutil
import os
import yaml
import requests
import streamlit as st
from typing import List, Dict, Optional, Tuple
from supabase import create_client, Client

import PyPDF2
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

# ── 数据库配置 (优先尝试 Streamlit Secrets，否则尝试环境变量) ───────────────
try:
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    supabase: Optional[Client] = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception:
    supabase = None # 如果没有配置，将回退到本地模式或报错反馈 (此处可根据需要微调)

# 如果 supabase 不可用，本地测试依然使用 styles 目录
STYLE_DIR = "styles"
if not os.path.exists(STYLE_DIR):
    os.makedirs(STYLE_DIR)

# ── 路径配置 ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
STYLES_DIR = BASE_DIR / "styles"
SAMPLES_DIR = BASE_DIR / "samples"
OUTPUT_DIR = BASE_DIR / "output"

for _d in [STYLES_DIR, SAMPLES_DIR, OUTPUT_DIR]:
    _d.mkdir(parents=True, exist_ok=True)

# ── 模型配置表 ────────────────────────────────────────────────────────────────
PROVIDERS = {
    "🇨🇳 DeepSeek": {
        "base_url": "https://api.deepseek.com",
        "models": ["deepseek-chat", "deepseek-reasoner"],
        "type": "openai_compatible",
    },
    "🇺🇸 OpenAI": {
        "base_url": "https://api.openai.com/v1",
        "models": ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"],
        "type": "openai_compatible",
    },
    "🇺🇸 Claude (Anthropic)": {
        "base_url": None,
        "models": ["claude-opus-4-5", "claude-sonnet-4-5", "claude-3-5-haiku-20241022"],
        "type": "anthropic",
    },
    "🌐 Google Gemini": {
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
        "models": ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash"],
        "type": "openai_compatible",
    },
    "🇨🇳 Moonshot (Kimi)": {
        "base_url": "https://api.moonshot.cn/v1",
        "models": ["moonshot-v1-8k", "moonshot-v1-32k", "moonshot-v1-128k"],
        "type": "openai_compatible",
    },
    "🇨🇳 智谱 GLM": {
        "base_url": "https://open.bigmodel.cn/api/paas/v4/",
        "models": ["glm-4", "glm-4-flash", "glm-4-plus"],
        "type": "openai_compatible",
    },
    "🇨🇳 通义千问": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "models": ["qwen-max", "qwen-plus", "qwen-turbo"],
        "type": "openai_compatible",
    },
}

# ── AI 调用 ───────────────────────────────────────────────────────────────────
def call_ai(prompt: str, api_key: str, provider_name: str, model_name: str) -> str:
    provider = PROVIDERS[provider_name]
    if provider["type"] == "openai_compatible":
        client = OpenAI(api_key=api_key, base_url=provider["base_url"])
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=8000,
        )
        return response.choices[0].message.content or ""
    elif provider["type"] == "anthropic":
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model=model_name,
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}],
            )
            return response.content[0].text
        except ImportError:
            return "❌ 使用 Claude 需要先安装 anthropic 库：pip install anthropic"
    return "❌ 未知的模型类型"

# ── 风格操作 (改为 Supabase 云端存储) ──────────────────────────────────────────

def get_style_tree() -> Dict[str, List[str]]:
    """从数据库获取风格树结构 {分类: [风格1, 风格2]}"""
    if not supabase:
        return _get_local_style_tree()

    try:
        data = supabase.table("styles").select("category, style_name").execute()
        tree = {}
        for row in data.data:
            cat = row["category"]
            style = row["style_name"]
            if cat not in tree:
                tree[cat] = []
            tree[cat].append(style)
        
        # Sort styles within each category and then sort categories
        sorted_tree = {k: sorted(v) for k, v in sorted(tree.items())}
        
        # Ensure "通用" is at the beginning of its category if it exists
        for cat, styles in sorted_tree.items():
            if "通用" in styles:
                styles.remove("通用")
                sorted_tree[cat] = ["通用"] + styles
        
        return sorted_tree
    except Exception as e:
        st.warning(f"无法从云端获取风格树: {e}，正在尝试本地备选...")
        return _get_local_style_tree()


def get_categories() -> list[str]:
    return list(get_style_tree().keys())

def get_styles_in_category(category: str) -> list[str]:
    tree = get_style_tree()
    styles = tree.get(category, [])
    # The get_style_tree already handles sorting and "通用" placement,
    # so we just return the list directly.
    return styles

def read_style_file(category: str, style_name: str) -> Optional[str]:
    """从数据库读取具体风格内容"""
    if not supabase:
        return _read_local_style_file(category, style_name)

    try:
        data = supabase.table("styles").select("content")\
            .match({"category": category, "style_name": style_name})\
            .maybe_single().execute()
        return data.data["content"] if data.data else None
    except Exception as e:
        st.warning(f"云端读取失败: {e}")
        return _read_local_style_file(category, style_name)


def save_style_file(category: str, style_name: str, content: str) -> Path:
    """保存或更新数据库中的风格 (Upsert)"""
    if not supabase:
        _save_local_style_file(category, style_name, content)
        return STYLES_DIR / category / f"{style_name}.md" # Return local path for consistency

    try:
        # 使用 upsert，如果分类+风格名已存在则更新，不存在则插入
        supabase.table("styles").upsert({
            "category": category,
            "style_name": style_name,
            "content": content,
            "updated_at": "now()"
        }, on_conflict="category,style_name").execute()
        return Path(f"supabase://styles/{category}/{style_name}.md") # Placeholder for Supabase path
    except Exception as e:
        st.error(f"云端保存失败: {e}")
        _save_local_style_file(category, style_name, content)
        return STYLES_DIR / category / f"{style_name}.md" # Return local path on failure


def delete_style_file(category: str, style_name: str):
    """从数据库删除指定风格"""
    if not supabase:
        _delete_local_style_file(category, style_name)
        return

    try:
        supabase.table("styles").delete()\
            .match({"category": category, "style_name": style_name})\
            .execute()
    except Exception as e:
        st.error(f"云端删除失败: {e}")


def rename_style(old_cat: str, old_name: str, new_cat: str, new_name: str):
    """重命名或移动数据库中的风格记录"""
    if not supabase:
        _rename_local_style(old_cat, old_name, new_cat, new_name)
        return

    try:
        supabase.table("styles").update({
            "category": new_cat,
            "style_name": new_name,
            "updated_at": "now()"
        }).match({"category": old_cat, "style_name": old_name})\
        .execute()
    except Exception as e:
        st.error(f"云端重命名失败: {e}")


# ── 本地存储备选逻辑 (供无网络或未配置时回退使用) ───────────────────────────

def _get_local_style_tree():
    """返回 {category: [style_names]} 有序字典，通用 排在第一位"""
    tree: dict[str, list[str]] = {}
    if not STYLES_DIR.exists():
        return tree
    for item in sorted(STYLES_DIR.iterdir()):
        if item.is_dir() and not item.name.startswith("."):
            styles = sorted([f.stem for f in item.glob("*.md")])
            if styles:
                # 把"通用"排到最前面
                if "通用" in styles:
                    styles = ["通用"] + [s for s in styles if s != "通用"]
                tree[item.name] = styles
    return tree

def _read_local_style_file(category: str, style_name: str) -> str:
    path = STYLES_DIR / category / f"{style_name}.md"
    return path.read_text(encoding="utf-8") if path.exists() else ""

def _save_local_style_file(category: str, style_name: str, content: str):
    cat_dir = STYLES_DIR / category
    cat_dir.mkdir(parents=True, exist_ok=True)
    path = cat_dir / f"{style_name}.md"
    path.write_text(content, encoding="utf-8")

def _delete_local_style_file(category: str, style_name: str):
    path = STYLES_DIR / category / f"{style_name}.md"
    if path.exists():
        path.unlink()
    cat_dir = STYLES_DIR / category
    if cat_dir.exists() and not any(cat_dir.glob("*.md")):
        cat_dir.rmdir()

def _rename_local_style(old_cat: str, old_name: str, new_cat: str, new_name: str):
    content = _read_local_style_file(old_cat, old_name)
    _save_local_style_file(new_cat, new_name, content)
    _delete_local_style_file(old_cat, old_name)

# ── AI 自动命名 ───────────────────────────────────────────────────────────────
def ai_suggest_category(text: str, api_key: str, provider_name: str, model_name: str) -> str:
    prompt = (
        "根据以下政府文件的内容，判断它最适合归入哪类文件类型。\n"
        "常见类型：政策行动方案、工作报告、通知公告、调研报告、规划纲要、"
        "工作总结、意见建议、实施方案、管理办法等。\n"
        "你也可以根据内容创建准确的新类型名称，长度4-10个字。\n"
        "只输出类别名称，不要任何解释和标点符号。\n\n"
        f"文件内容节选：\n{text[:2000]}"
    )
    result = call_ai(prompt, api_key, provider_name, model_name)
    return result.strip().replace("。", "").replace("，", "")[:20]

def ai_suggest_style_name(text: str, api_key: str, provider_name: str, model_name: str) -> str:
    prompt = (
        "根据以下政府文件的写作风格和内容主题，为这个文风模板起一个简洁的中文名称。\n"
        "名称应体现文件的行业领域或具体主题（如：生态环境保护、能源双碳、乡村振兴、数字经济等），"
        "长度4-10个字。如果无法判断具体领域，则描述文风特点（如：简洁务实型、正式公告型等）。\n"
        "只输出名称，不要任何解释和标点符号。\n\n"
        f"文件内容节选：\n{text[:2000]}"
    )
    result = call_ai(prompt, api_key, provider_name, model_name)
    return result.strip().replace("。", "").replace("，", "")[:20]

# ── 去 AI 化（基于 Humanizer-zh 规则）────────────────────────────────────────
HUMANIZER_PROMPT = """\
你是一名经验丰富的政府文件编辑，对 AI 生成的文稿进行「去 AI 化」处理，使其读起来更自然、更像资深人工撰写。

请对照以下 24 种 AI 写作痕迹逐一检查并改写：

【内容模式】
1. 过度强调"意义/格局/深远影响/历史性"等空洞词——改为具体描述
2. 用"专家认为/观察者指出/研究表明"等模糊归因——改为具体来源或删除
3. 公式化的"挑战与展望"段落——如无实质内容则删除或改写为具体措施
4. 宣传广告式语言（"充分体现/深刻彰显/有力推动"）——改为平实表达

【语言和语法模式】
5. 高频 AI 词汇替换：
   此外→另外|至关重要→关键|深入探讨→分析|增强→提高|培养→建立|
   彰显→体现|错综复杂→复杂|宝贵的→重要的|充满活力的→活跃的|
   深刻→（直接删除或换具体词）|持久的→长期的
6. 否定式排比泛滥（"不仅…而且…更…"）——拆解为独立句子
7. 三段式列举过度（"A、B 与 C"）——改为两项或具体描述
8. 系动词回避（过度用分词结构代替"是"）——恢复直接表达
9. 刻意换词导致同义词循环——选定一个词一贯使用
10. "虚假范围"（"在许多方面""在很大程度上"）——具体化或删除

【风格模式】
11. 破折号过度使用——改用句号或逗号
12. 粗体滥用——只保留真正关键的标注
13. 每段都是"**标题**：说明"格式——改用自然段落叙述
14. 去除所有表情符号

【交流模式和填充词】
15. 协作痕迹（"如您所要求的/根据您的问题"）——直接删除
16. 谄媚语气（"这是个很好的问题/非常重要的议题"）——直接删除
17. 填充短语："值得注意的是/不可否认的是/毋庸置疑/有目共睹"——删除或保留后续实质
18. 过度限定滥用："可能/相关/一定程度上/某种意义上"每句都有——减少到必要处
19. 通用积极结论（"相信必将取得更大成效/为事业做出新贡献"）——删除或改为具体预期
20. 开头填充词（"当然/好的/明白了/当谈到…时"）——直接删除
21. "知识截止日期免责声明"——删除
22. 无实质的强调："至关重要的一点是""需要特别指出的是"——删除或直接陈述
23. 完美的三段论结构——打破机械感，允许段落长短不一
24. "金句化"结尾（听起来像名言警句的句子）——改为具体、实在的陈述

处理原则：
✓ 保留全部核心信息、政策数据和逻辑框架，不改变事实
✓ 维持政府公文的正式性和权威感，不要改成散文
✓ 让文字更具体、直接，有实质内容
✓ 句子长短适当混搭，避免全是长句或全是短句

直接输出处理后的完整文稿，不要解释你修改了什么。

【待处理文稿】
"""

def humanize_text(text: str, api_key: str, provider_name: str, model_name: str) -> str:
    return call_ai(HUMANIZER_PROMPT + text, api_key, provider_name, model_name)

# ── 文件读取 ──────────────────────────────────────────────────────────────────
def read_uploaded_file(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix in [".txt", ".md"]:
        return uploaded_file.read().decode("utf-8")
    elif suffix == ".pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    elif suffix == ".docx":
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    return "不支持的格式"

# ── 风格提取 ──────────────────────────────────────────────────────────────────
def extract_style_content(combined_text: str, style_name: str,
                           api_key: str, provider_name: str, model_name: str) -> str:
    prompt = f"""\
你是一名语言风格分析专家。请分析以下文件的语言风格，只分析文字表达风格，不分析文件结构。

【文件内容】
{combined_text[:6000]}

请按以下维度输出风格分析报告（Markdown格式）：

# {style_name} 文风模板

## 1. 整体语言风格
（用3-5句话描述整体语言气质：正式程度、客观程度、语气特点等）

## 2. 常用动词和搭配
（列出最高频的动词及固定搭配，例如："推进……""加快……""积极……"）

## 3. 句式特点
（描述句子典型特征：长短、是否多排比、是否多被动句等）

## 4. 常用连接词和过渡语
（列出段落/句子间常用连接词，例如："为此""对此""与此同时"）

## 5. 用词偏好
（描述词汇选择特点：偏正式/口语、专业术语多寡、是否避免某些词汇等）

## 6. 语气和态度
（描述整体语气：指令性、建议性、还是陈述性）

## 7. 典型例句
（从原文摘录10个最能代表这种文风的句子，只选正文句子，不选标题）
"""
    return call_ai(prompt, api_key, provider_name, model_name)


def extract_general_style(combined_text: str, category_name: str,
                           api_key: str, provider_name: str, model_name: str) -> str:
    """
    提炼分类级别的「通用」写作风格：
    聚焦跨主题的共性语言规律，不依赖具体主题内容。
    """
    prompt = f"""\
你是一名语言风格分析专家。以下是「{category_name}」类政府文件的写作样本。
请提炼这类文件跨主题、跨内容的**通用语言风格规律**，
不要分析具体主题内容，只关注语言表达的共性特征。

【文件内容】
{combined_text[:6000]}

请按以下维度输出通用风格分析报告（Markdown格式）：

# 【{category_name}】通用文风模板

> 本模板提炼该类文件的**通用语言风格**，适用于所有「{category_name}」类文件的风格迁移基线。

## 1. 整体语言气质
（描述这类文件区别于其他文件类型的整体语言特点：正式程度、权威感、客观程度、语气等）

## 2. 高频动词与固定搭配
（列出这类文件中普遍高频出现的动词和固定搭配，不依赖具体主题）

## 3. 典型句式结构
（描述这类文件共有的句式偏好：是否多长句、排比、条件句、被动句等）

## 4. 段落组织逻辑
（描述段落内部和段落间的常见逻辑结构：总-分、时间顺序、问题-措施等）

## 5. 通用连接词与过渡语
（列出这类文件中跨主题都会用到的连接词、固定表达）

## 6. 用词风格与禁用表达
（描述词汇选择的总体风格，以及这类文件应避免的口语、俗语等表达）

## 7. 典型例句（通用性强）
（从原文摘录8-10个能够体现这类文件普遍风格、与具体主题无关的典型句子）
"""
    return call_ai(prompt, api_key, provider_name, model_name)

# ── 生成初稿 ──────────────────────────────────────────────────────────────────
def generate_draft(category: str, style_name: str, topic: str, key_points: str,
                   word_count: int, api_key: str, provider_name: str, model_name: str) -> str:
    style_template = read_style_file(category, style_name)
    prompt = f"""\
你是浙江省政府智库的资深政策研究员。请严格参照以下文风模板的语言风格来写作。

【文风模板】
{style_template}

请撰写一份文件：
- 主题：{topic}
- 字数：约{word_count}字
- 核心要点：
{key_points}

写作规范：
1. 语言风格必须严格参照上方文风模板
2. 数据暂时用[待填写]占位
3. 直接输出文件内容，不加解释
"""
    result = call_ai(prompt, api_key, provider_name, model_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    (OUTPUT_DIR / f"{topic[:15]}_{ts}.md").write_text(result, encoding="utf-8")
    return result

# ── 润色草稿 ──────────────────────────────────────────────────────────────────
def polish_draft(draft: str, category: str, style_name: str,
                 api_key: str, provider_name: str, model_name: str) -> str:
    style_template = read_style_file(category, style_name)
    prompt = f"""\
你是政府文件润色专家。请参照以下文风模板，对草稿进行润色：

【文风模板】
{style_template}

润色要求：
1. 保留核心信息和逻辑不变
2. 口语化表达改为正式公文用语
3. 语言风格向文风模板靠拢
4. 数据和事实不要改动

【待润色草稿】
{draft}

直接输出润色后的文件，不加解释。
"""
    result = call_ai(prompt, api_key, provider_name, model_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    (OUTPUT_DIR / f"润色_{ts}.md").write_text(result, encoding="utf-8")
    return result

# ── Markdown → Word 转换 ──────────────────────────────────────────────────────
def _add_bold_run(paragraph, text: str, base_font: str = "仿宋_GB2312", base_size: int = 12):
    """将含 **粗体** 标记的文本分段添加到段落"""
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.name = base_font
        run.font.size = Pt(base_size)


def markdown_to_docx(markdown_text: str, doc_title: str = "文档") -> io.BytesIO:
    doc = Document()

    # 页面设置（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    lines = markdown_text.split("\n")
    for line in lines:
        line = line.rstrip()

        if line.startswith("# "):
            h = doc.add_heading("", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)

        elif line.startswith("## "):
            h = doc.add_heading("", level=2)
            run = h.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)

        elif line.startswith("### "):
            h = doc.add_heading("", level=3)
            run = h.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(12)

        elif line.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.bold = True
            run.font.size = Pt(12)

        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_bold_run(p, line[2:].strip())

        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\.\s", "", line).strip()
            _add_bold_run(p, content)

        elif line.startswith(("---", "___", "===")):
            doc.add_paragraph("─" * 40)

        elif line:
            p = doc.add_paragraph()
            _add_bold_run(p, line)

        else:
            # 空行：仅当前一个段落有内容时才插入
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    buf.close()
    return docx_bytes
